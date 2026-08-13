import os
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

# Create directories
sample_dir = Path("sample_resumes")
sample_dir.mkdir(exist_ok=True)

def create_pdf_resume(filepath, name, email, education, experience, projects, skills):
    doc = SimpleDocTemplate(str(filepath), pagesize=letter)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=20, leading=24)
    sec_style = ParagraphStyle('Sec', parent=styles['Heading2'], fontSize=14, leading=18, spaceBefore=10)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14)
    
    story = []
    story.append(Paragraph(name, title_style))
    story.append(Paragraph(f"Email: {email}", body_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Education", sec_style))
    story.append(Paragraph(education, body_style))
    
    story.append(Paragraph("Experience", sec_style))
    story.append(Paragraph(experience, body_style))
    
    story.append(Paragraph("Projects", sec_style))
    story.append(Paragraph(projects, body_style))
    
    story.append(Paragraph("Skills", sec_style))
    story.append(Paragraph(skills, body_style))
    
    doc.build(story)
    print(f"Created PDF resume: {filepath}")

def create_docx_resume(filepath, name, email, education, experience, projects, skills):
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email}")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)
    
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(experience)
    
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(projects)
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)
    
    doc.save(str(filepath))
    print(f"Created DOCX resume: {filepath}")

if __name__ == "__main__":
    # 1. Data Analyst Resume
    create_pdf_resume(
        sample_dir / "Data_Analyst_Resume.pdf",
        "Jane Doe",
        "jane.doe@example.com",
        "B.S. in Statistics, State University, 2024",
        "Junior Data Analyst at RetailCorp (2024-2026)\nAnalyzed sales datasets, built reports, and optimized queries.",
        "Sales Performance Dashboard using Tableau\nCustomer Churn Clustering in Jupyter Notebook using Pandas and NumPy",
        "SQL, Python, Pandas, NumPy, Tableau, PowerBI, Excel, Communication, Statistics"
    )
    
    # 2. ML Engineer Resume
    create_docx_resume(
        sample_dir / "ML_Engineer_Resume.docx",
        "John Smith",
        "john.smith@example.com",
        "M.S. in Computer Science, Tech Institute, 2023",
        "Machine Learning Intern at AI Lab (2023-2025)\nDeveloped CNN models, optimized hyperparameters, and cleaned dataset tokens.",
        "Predictive Maintenance API using FastAPI and Scikit-learn\nReal-time object classifier using PyTorch and OpenCV",
        "Python, Scikit-learn, TensorFlow, PyTorch, NumPy, Pandas, Git, Docker, Machine Learning, Regression, Classification"
    )
    
    # 3. Cloud Engineer Resume
    create_pdf_resume(
        sample_dir / "Cloud_Engineer_Resume.pdf",
        "Bob Johnson",
        "bob.johnson@example.com",
        "B.Tech in Information Technology, Poly University, 2022",
        "Cloud Associate at TechCorp (2022-2026)\nManaged AWS resources, wrote bash scripts, and automated infrastructure deployments.",
        "Automated AWS Deployments using Terraform\nConfigured autoscaling Kubernetes clusters with Helm charts",
        "AWS, GCP, Azure, Docker, Kubernetes, Terraform, Ansible, Linux, Bash, CI/CD, Network Security"
    )
